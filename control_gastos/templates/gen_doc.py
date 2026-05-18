#!/usr/bin/env python3
"""Genera el documento Word de documentación técnica de FinTrack."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Estilos base ──────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1C, 0x24, 0x63)
DARK   = RGBColor(0x1A, 0x1F, 0x35)
GRAY   = RGBColor(0x6B, 0x72, 0x80)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREEN  = RGBColor(0x2E, 0x7D, 0x32)
RED    = RGBColor(0xC6, 0x28, 0x28)
LNAVY  = RGBColor(0xEE, 0xF0, 0xFA)

# Márgenes del documento
from docx.shared import Cm
sections = doc.sections
for section in sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Utilidades ────────────────────────────────────────────────────────────────
def add_heading(text, level=1, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20 if level == 1 else 15 if level == 2 else 13)
    run.font.color.rgb = color or NAVY
    run.font.name = 'Calibri'
    return p

def add_body(text, indent=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = color or DARK
    run.font.name = 'Calibri'
    return p

def add_bullet(text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.8)
    if bold_part:
        r1 = p.add_run(bold_part + ': ')
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = 'Calibri'
        r1.font.color.rgb = NAVY
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        r2.font.name = 'Calibri'
        r2.font.color.rgb = DARK
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.font.color.rgb = DARK
    return p

def add_code_block(lines):
    """Agrega un bloque de código con fondo gris claro."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)
        # Fondo gris claro via shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F2F5')
        pPr.append(shd)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1C, 0x24, 0x63)

def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1C2463')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════
# Bloque azul de título usando una tabla de una celda
table_cover = doc.add_table(rows=1, cols=1)
table_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table_cover.cell(0, 0)
shade_cell(cell, '1C2463')
cell.width = Inches(6)

p_logo = cell.paragraphs[0]
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_logo.paragraph_format.space_before = Pt(30)
r = p_logo.add_run('📊 FinTrack')
r.font.size  = Pt(32)
r.font.bold  = True
r.font.color.rgb = WHITE
r.font.name  = 'Calibri'

p_sub = cell.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run('Aplicación de Finanzas Personales')
r2.font.size  = Pt(15)
r2.font.color.rgb = RGBColor(0xCC, 0xD0, 0xF0)
r2.font.name  = 'Calibri'

p_ver = cell.add_paragraph()
p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ver.paragraph_format.space_after = Pt(30)
r3 = p_ver.add_run('Documentación Técnica — Versión 2.0 (Reingeniería Educativa)')
r3.font.size  = Pt(10)
r3.font.color.rgb = RGBColor(0xAA, 0xB0, 0xD8)
r3.font.name  = 'Calibri'

doc.add_paragraph()  # Espacio

# Metadatos de portada
meta = [
    ('Institución',  'Programación Estructurada y Funcional'),
    ('Docente',      '—'),
    ('Estudiante',   '—'),
    ('Fecha',        'Mayo 2026'),
    ('Herramientas', 'Python 3 · VS Code · Chart.js · reportlab'),
    ('IA de apoyo',  'Claude (Anthropic) — optimización y documentación'),
]
t_meta = doc.add_table(rows=len(meta), cols=2)
t_meta.style = 'Table Grid'
for i, (lbl, val) in enumerate(meta):
    c0 = t_meta.cell(i, 0)
    c1 = t_meta.cell(i, 1)
    shade_cell(c0, 'EEF0FA')
    shade_cell(c1, 'FFFFFF')
    r0 = c0.paragraphs[0].add_run(lbl)
    r0.bold = True; r0.font.size = Pt(10); r0.font.name = 'Calibri'; r0.font.color.rgb = NAVY
    r1 = c1.paragraphs[0].add_run(val)
    r1.font.size = Pt(10); r1.font.name = 'Calibri'; r1.font.color.rgb = DARK

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECCIÓN 2 – DESCRIPCIÓN DEL PROGRAMA
# ═══════════════════════════════════════════════════════════════
add_heading('2. Descripción del Programa')
add_divider()

add_heading('¿Qué hace el programa?', level=2)
add_body(
    'FinTrack es una aplicación web de finanzas personales que corre completamente en el '
    'navegador del usuario. Permite registrar gastos e ingresos con categoría, fecha y cuenta, '
    'visualizar el saldo disponible en tiempo real, establecer metas de ahorro con barra de '
    'progreso, analizar la distribución del gasto mediante una gráfica tipo donut, y exportar '
    'las transacciones en formato CSV o PDF (este último a través del servidor Python).'
)

add_heading('¿Cuál es su finalidad?', level=2)
add_body(
    'Proporcionar una herramienta didáctica y funcional que demuestre cómo construir una '
    'aplicación de gestión financiera personal aplicando únicamente paradigmas de programación '
    'estructurada y funcional básica, sin depender de frameworks ni programación orientada a objetos.'
)

add_heading('¿Qué problema busca resolver?', level=2)
add_body(
    'Muchas personas no tienen un control claro de sus finanzas personales porque las '
    'herramientas disponibles son complejas o requieren conexión a servicios en la nube. '
    'FinTrack resuelve esto con una aplicación local, ligera y fácil de usar que funciona '
    'offline, sin cuentas ni registro.'
)

add_heading('Requerimientos Funcionales', level=2)
req_func = [
    ('RF-01', 'Registrar transacciones con tipo (gasto/ingreso), monto, descripción, categoría, fecha y cuenta.'),
    ('RF-02', 'Calcular y mostrar el saldo total en tiempo real: saldo_inicial + ingresos - gastos.'),
    ('RF-03', 'Eliminar transacciones individuales con confirmación del usuario.'),
    ('RF-04', 'Buscar transacciones por descripción o categoría en tiempo real.'),
    ('RF-05', 'Gestionar metas de ahorro: nombre, monto objetivo, monto ahorrado y barra de progreso.'),
    ('RF-06', 'Mostrar gráfica donut con distribución porcentual de gastos por categoría.'),
    ('RF-07', 'Exportar transacciones en formato CSV descargable.'),
    ('RF-08', 'Exportar informe PDF completo vía servidor Python (requiere reportlab).'),
    ('RF-09', 'Persistir el estado completo en localStorage entre sesiones del navegador.'),
    ('RF-10', 'Mostrar insight automático: balance del período y categoría con mayor gasto.'),
]
t_rf = doc.add_table(rows=len(req_func)+1, cols=2)
t_rf.style = 'Table Grid'
# Cabecera
shade_cell(t_rf.cell(0, 0), '1C2463')
shade_cell(t_rf.cell(0, 1), '1C2463')
for col, txt in enumerate(['Código', 'Descripción del Requerimiento']):
    r = t_rf.cell(0, col).paragraphs[0].add_run(txt)
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE; r.font.name = 'Calibri'
for i, (cod, desc) in enumerate(req_func):
    fill = 'F8F9FB' if i % 2 == 0 else 'FFFFFF'
    shade_cell(t_rf.cell(i+1, 0), 'EEF0FA')
    shade_cell(t_rf.cell(i+1, 1), fill)
    rc = t_rf.cell(i+1, 0).paragraphs[0].add_run(cod)
    rc.bold = True; rc.font.size = Pt(10); rc.font.name = 'Calibri'; rc.font.color.rgb = NAVY
    rd = t_rf.cell(i+1, 1).paragraphs[0].add_run(desc)
    rd.font.size = Pt(10); rd.font.name = 'Calibri'; rd.font.color.rgb = DARK

doc.add_paragraph()
add_heading('Requerimientos No Funcionales', level=2)
req_nf = [
    ('RNF-01', 'Legibilidad',    'Código comentado bloque por bloque; nombres de función en español descriptivo.'),
    ('RNF-02', 'Mantenibilidad', 'Funciones cortas con responsabilidad única; estado global explícito.'),
    ('RNF-03', 'Rendimiento',    'Operaciones de renderizado O(n); sin dependencias pesadas en el cliente.'),
    ('RNF-04', 'Portabilidad',   'Funciona en cualquier navegador moderno; servidor requiere solo Python 3.8+.'),
    ('RNF-05', 'Seguridad',      'Escapado HTML en todas las salidas de usuario para prevenir XSS.'),
    ('RNF-06', 'Accesibilidad',  'Diseño responsivo con breakpoints en 600 px y 900 px.'),
]
t_nf = doc.add_table(rows=len(req_nf)+1, cols=3)
t_nf.style = 'Table Grid'
shade_cell(t_nf.cell(0, 0), '1C2463')
shade_cell(t_nf.cell(0, 1), '1C2463')
shade_cell(t_nf.cell(0, 2), '1C2463')
for col, txt in enumerate(['Código', 'Atributo', 'Descripción']):
    r = t_nf.cell(0, col).paragraphs[0].add_run(txt)
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE; r.font.name = 'Calibri'
for i, (cod, attr, desc) in enumerate(req_nf):
    fill = 'F8F9FB' if i % 2 == 0 else 'FFFFFF'
    for col in range(3): shade_cell(t_nf.cell(i+1, col), fill)
    shade_cell(t_nf.cell(i+1, 0), 'EEF0FA')
    vals = [cod, attr, desc]
    for col, v in enumerate(vals):
        r = t_nf.cell(i+1, col).paragraphs[0].add_run(v)
        r.font.size = Pt(10); r.font.name = 'Calibri'; r.font.color.rgb = DARK
        if col < 2: r.bold = True; r.font.color.rgb = NAVY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECCIÓN 3 – PROBLEMA A RESOLVER
# ═══════════════════════════════════════════════════════════════
add_heading('3. Problema a Resolver')
add_divider()

add_body(
    'Una persona recibe su salario mensual y a mitad de mes no sabe en qué gastó su dinero '
    'ni si podrá alcanzar su meta de ahorro para las vacaciones. Necesita una herramienta simple '
    'que le permita registrar cada gasto o ingreso, ver cuánto le queda disponible y '
    'saber qué categoría le está costando más dinero.'
)

doc.add_paragraph()
add_heading('Definición Formal: Entradas, Proceso y Salidas', level=2)

eps = [
    ('ENTRADAS', '1C2463', 'FFFFFF', [
        'Saldo inicial definido por el usuario (número decimal positivo).',
        'Transacciones: tipo (gasto/ingreso), monto, descripción, categoría, fecha y cuenta.',
        'Metas de ahorro: nombre, monto objetivo y monto ya ahorrado.',
        'Consulta de búsqueda (texto libre para filtrar transacciones).',
    ]),
    ('PROCESO', '2E7D32', 'FFFFFF', [
        'Almacenar transacciones en un array en memoria y persistir en localStorage.',
        'Calcular saldo_actual = saldo_inicial + Σ(ingresos) - Σ(gastos) con bucle for.',
        'Agrupar gastos por categoría usando arrays paralelos para el gráfico.',
        'Filtrar transacciones según el texto de búsqueda con comparación de cadenas.',
        'Ordenar transacciones por fecha con algoritmo de selección (burbuja).',
        'Calcular porcentaje de progreso de cada meta: (ahorrado / objetivo) × 100.',
        'Generar PDF en el servidor Python con tablas de resumen y detalle.',
    ]),
    ('SALIDAS', 'C62828', 'FFFFFF', [
        'Saldo actual visible en tiempo real en la tarjeta principal.',
        'Lista de transacciones ordenadas con categoría, monto y botón de eliminar.',
        'Gráfica donut con distribución porcentual de gastos por categoría.',
        'Barras de progreso de cada meta de ahorro con porcentaje completado.',
        'Mensaje de insight: balance del período y categoría de mayor gasto.',
        'Archivo CSV descargable con todas las transacciones.',
        'Archivo PDF descargable con resumen financiero y listado (vía servidor).',
    ]),
]

for (titulo, fill_hex, txt_hex, items) in eps:
    t = doc.add_table(rows=1+len(items), cols=1)
    t.style = 'Table Grid'
    # Cabecera de color
    hc = t.cell(0, 0)
    shade_cell(hc, fill_hex)
    rh = hc.paragraphs[0].add_run(titulo)
    rh.bold = True; rh.font.size = Pt(11); rh.font.name = 'Calibri'
    rh.font.color.rgb = WHITE
    # Ítems
    for i, item in enumerate(items):
        c = t.cell(i+1, 0)
        shade_cell(c, 'F8F9FB' if i % 2 == 0 else 'FFFFFF')
        ri = c.paragraphs[0].add_run('• ' + item)
        ri.font.size = Pt(10); ri.font.name = 'Calibri'; ri.font.color.rgb = DARK
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECCIÓN 4 – HERRAMIENTAS UTILIZADAS
# ═══════════════════════════════════════════════════════════════
add_heading('4. Herramientas Utilizadas')
add_divider()

herramientas = [
    ('🐍', 'Lenguaje: Python 3.8+',
     'Utilizado para el servidor HTTP local (fintrack_server.py). '
     'Maneja las rutas GET/POST, genera los informes PDF y sirve el archivo HTML '
     'estático. Se usan únicamente las librerías estándar http.server, socketserver, '
     'json, io, sys y pathlib, más la librería opcional reportlab para PDF.'),

    ('🌐', 'Lenguaje: HTML5 + CSS3 + JavaScript (ES6)',
     'El frontend completo de FinTrack se implementa en un único archivo HTML. '
     'El JavaScript sigue el paradigma estructurado y funcional: solo variables, '
     'arrays, funciones, if/else y bucles for/while. Sin clases ni diccionarios complejos.'),

    ('💻', 'IDE: Visual Studio Code',
     'Editor de código principal. Se recomienda instalar las extensiones '
     '"Prettier" (formateo automático), "Python" (depuración), y '
     '"Live Server" (recarga automática del HTML).'),

    ('📊', 'Librería de gráficas: Chart.js 4.4.1 (CDN)',
     'Se usa únicamente para la visualización del gráfico donut en la vista Analytics. '
     'Los datos que alimentan la gráfica (agrupación, sumas, porcentajes) se calculan '
     'manualmente con bucles for en JavaScript, sin delegar la lógica a la librería.'),

    ('📄', 'Generación de PDF: reportlab (opcional)',
     'Librería Python para construir documentos PDF con tablas, estilos y párrafos. '
     'Se instala con: pip install reportlab. Si no está instalada, el servidor '
     'continúa funcionando para todas las demás funciones.'),

    ('🤖', 'Inteligencia Artificial: Claude (Anthropic)',
     'Utilizado como herramienta de apoyo para: revisión de la lógica de los algoritmos, '
     'optimización de la estructura del código, generación de comentarios pedagógicos '
     'y redacción de esta documentación técnica. El diseño y la implementación '
     'son responsabilidad del equipo de desarrollo.'),

    ('📐', 'Pseudocódigo: PSeInt / Flowgorithm',
     'Herramientas recomendadas para planificar y validar los algoritmos principales '
     '(cálculo de saldo, filtrado, ordenamiento) antes de implementarlos. '
     'Permiten visualizar el flujo de control sin preocuparse por la sintaxis.'),
]

for (icono, titulo, desc) in herramientas:
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.columns[0].width = Cm(1.5)
    c_icon = t.cell(0, 0)
    c_desc = t.cell(0, 1)
    shade_cell(c_icon, 'EEF0FA')
    shade_cell(c_desc, 'FFFFFF')
    ri = c_icon.paragraphs[0].add_run(icono)
    ri.font.size = Pt(20)
    ri.font.name = 'Calibri'
    c_icon.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = c_desc.paragraphs[0].add_run(titulo)
    rt.bold = True; rt.font.size = Pt(11); rt.font.name = 'Calibri'; rt.font.color.rgb = NAVY
    rd = c_desc.add_paragraph().add_run(desc)
    rd.font.size = Pt(10); rd.font.name = 'Calibri'; rd.font.color.rgb = DARK
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECCIÓN 5 – CÓDIGO FUENTE (referencia y estructura)
# ═══════════════════════════════════════════════════════════════
add_heading('5. Código Fuente — Estructura y Secciones')
add_divider()

add_body(
    'Ambos archivos están disponibles como entregables separados (fintrack.html y fintrack_server.py). '
    'A continuación se documenta la estructura modular de cada uno, con la descripción '
    'de cada sección y los fragmentos de código más representativos.'
)

# Sub: fintrack.html
add_heading('5.1  fintrack.html — Frontend (HTML + CSS + JavaScript)', level=2)

secciones_html = [
    ('S-01', 'DATOS CONSTANTES',
     'Arrays paralelos CAT_IDS, CAT_NOMBRES, CAT_ICONOS, CAT_COLORES y COLORES_GRAFICA.',
     'Definen las 20 categorías disponibles usando 4 arrays de igual longitud. '
     'Se accede por índice numérico en lugar de usar objetos (sin POO).'),
    ('S-02', 'ESTADO GLOBAL',
     'Variables: saldoInicial, transacciones[], metas[], tipoTxActual, categoriaSeleccionada.',
     'El estado completo de la app vive en variables globales simples. '
     'Las transacciones son arrays de 7 elementos: [id, tipo, monto, desc, cat, fecha, cuenta].'),
    ('S-03', 'PERSISTENCIA',
     'Funciones guardarEstado() y cargarEstado().',
     'Serializan/deserializan el estado con JSON.stringify/parse desde localStorage.'),
    ('S-04', 'FUNCIONES AUXILIARES',
     'obtenerIndiceCat(id), nombreCat(id), iconoCat(id), colorCat(id).',
     'Buscan información de categoría usando un bucle for sobre el array CAT_IDS.'),
    ('S-05', 'FUNCIONES DE FORMATO',
     'formatearEuro(n), formatearDolar(n), formatearFecha(iso), escaparHTML(str).',
     'Convierten tipos primitivos en cadenas con formato visual para el DOM.'),
    ('S-06', 'CÁLCULOS FINANCIEROS',
     'calcularSaldoReal(), calcularTotalGastos(), calcularTotalIngresos().',
     'Recorren el array de transacciones con bucles for y acumulan los totales.'),
    ('S-07', 'RENDERIZADO',
     'renderizar(), renderizarTransacciones(), renderizarResumen(), renderizarMetas(), renderizarInsight(), renderizarAnalytics().',
     'Funciones de actualización del DOM. Construyen HTML con concatenación de cadenas en bucles for.'),
    ('S-08', 'CONTROL DE VISTAS',
     'mostrarVista(nombre, botonEl), seleccionarPeriodo(periodo, botonEl).',
     'Alternan la visibilidad de #view-dashboard y #view-analytics.'),
    ('S-09', 'MODAL TRANSACCIÓN',
     'abrirModal(), cerrarModal(), establecerTipo(), renderizarGridCategorias(), seleccionarCat(), guardarTransaccion(), eliminarTransaccion().',
     'Gestión completa del formulario de registro. Validación con if/else.'),
    ('S-10', 'MODAL SALDO',
     'abrirModalSaldo(), cerrarModalSaldo(), guardarSaldo().',
     'Permite al usuario editar el saldo de partida.'),
    ('S-11', 'MODAL METAS',
     'abrirModalMeta(), cerrarModalMeta(), guardarMeta(), eliminarMeta().',
     'CRUD de metas de ahorro en el array metas[].'),
    ('S-12', 'BÚSQUEDA Y EXPORTAR',
     'filtrarTransacciones(), exportarCSV().',
     'Filtrado en tiempo real y generación de archivo CSV con bucle for.'),
    ('S-13', 'CIERRE DE MODALES', 'Listener sobre .modal-overlay.',
     'Cierra el modal al hacer clic en el fondo oscuro del overlay.'),
    ('S-14', 'INICIALIZACIÓN', 'IIFE inicializar().',
     'Punto de entrada: carga el estado, ofrece configurar el saldo y renderiza.'),
]

t_html = doc.add_table(rows=len(secciones_html)+1, cols=4)
t_html.style = 'Table Grid'
cabeceras = ['ID', 'Sección', 'Funciones', 'Descripción']
for col, txt in enumerate(cabeceras):
    shade_cell(t_html.cell(0, col), '1C2463')
    r = t_html.cell(0, col).paragraphs[0].add_run(txt)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE; r.font.name = 'Calibri'
for i, (sid, sec, funcs, desc) in enumerate(secciones_html):
    fill = 'F8F9FB' if i % 2 == 0 else 'FFFFFF'
    for col in range(4): shade_cell(t_html.cell(i+1, col), fill)
    shade_cell(t_html.cell(i+1, 0), 'EEF0FA')
    vals = [sid, sec, funcs, desc]
    for col, v in enumerate(vals):
        r = t_html.cell(i+1, col).paragraphs[0].add_run(v)
        r.font.size = Pt(9); r.font.name = 'Calibri'; r.font.color.rgb = DARK
        if col <= 1: r.bold = True; r.font.color.rgb = NAVY

doc.add_paragraph()

# Fragmento de código representativo (cálculo de saldo)
add_heading('Fragmento representativo — Cálculo del saldo real', level=2)
add_body(
    'La siguiente función ilustra el paradigma estructurado aplicado: '
    'usa solo una variable acumuladora y un bucle for con if/else.'
)
add_code_block([
    'function calcularSaldoReal() {',
    '    var total = saldoInicial;    // Partir del saldo base',
    '    for (var i = 0; i < transacciones.length; i++) {',
    '        var tx = transacciones[i];',
    '        if (tx[1] === "income") {',
    '            total += tx[2];      // Sumar ingresos',
    '        } else {',
    '            total -= tx[2];      // Restar gastos',
    '        }',
    '    }',
    '    return total;',
    '}',
])

doc.add_paragraph()

# Sub: fintrack_server.py
add_heading('5.2  fintrack_server.py — Backend (Python)', level=2)

secciones_py = [
    ('S-01', 'CONFIGURACIÓN',    'Variables PUERTO, ARCHIVO_HTML, DIRECTORIO.',
     'Constantes de configuración del servidor. Modificables sin tocar el resto del código.'),
    ('S-02', 'DEPENDENCIA PDF',  'Bloque try/except para reportlab.',
     'Detecta si reportlab está instalado. Si no, el servidor sigue funcionando sin PDF.'),
    ('S-03', 'LÓGICA FINANCIERA','calcular_totales(txs), ordenar_por_fecha(txs), validar_datos_pdf(datos).',
     'Funciones puras de negocio. Usan bucles for y if/else. Sin clases propias.'),
    ('S-04', 'GENERACIÓN PDF',   'construir_tabla_resumen(), construir_tabla_transacciones(), generar_pdf().',
     'Coordinan la construcción del PDF con reportlab en funciones de responsabilidad única.'),
    ('S-05', 'RESPUESTAS HTTP',  'enviar_respuesta_json(), manejar_estado(), manejar_exportar_pdf().',
     'Funciones que construyen y envían respuestas HTTP. Separadas del manejador.'),
    ('S-06', 'MANEJADOR HTTP',   'Clase ManejadorFinTrack (mínima herencia requerida).',
     'Subclase de SimpleHTTPRequestHandler. Solo enruta peticiones a las funciones de negocio.'),
    ('S-07', 'INICIALIZACIÓN',   'verificar_archivo_html(), mostrar_banner(), iniciar_servidor().',
     'Funciones del flujo principal. Se llaman en secuencia desde el bloque __main__.'),
]

t_py = doc.add_table(rows=len(secciones_py)+1, cols=4)
t_py.style = 'Table Grid'
for col, txt in enumerate(['ID', 'Sección', 'Funciones', 'Descripción']):
    shade_cell(t_py.cell(0, col), '1C2463')
    r = t_py.cell(0, col).paragraphs[0].add_run(txt)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE; r.font.name = 'Calibri'
for i, (sid, sec, funcs, desc) in enumerate(secciones_py):
    fill = 'F8F9FB' if i % 2 == 0 else 'FFFFFF'
    for col in range(4): shade_cell(t_py.cell(i+1, col), fill)
    shade_cell(t_py.cell(i+1, 0), 'EEF0FA')
    vals = [sid, sec, funcs, desc]
    for col, v in enumerate(vals):
        r = t_py.cell(i+1, col).paragraphs[0].add_run(v)
        r.font.size = Pt(9); r.font.name = 'Calibri'; r.font.color.rgb = DARK
        if col <= 1: r.bold = True; r.font.color.rgb = NAVY

doc.add_paragraph()
add_heading('Fragmento representativo — Ordenamiento por fecha (Python)', level=2)
add_body(
    'Se usa un algoritmo de selección explícito con for/for anidado '
    'en lugar de la función sorted() nativa, para evidenciar la lógica del proceso.'
)
add_code_block([
    'def ordenar_por_fecha(transacciones):',
    '    lista = transacciones[:]      # Copia para no modificar el original',
    '    n = len(lista)',
    '    for i in range(n - 1):',
    '        max_idx = i',
    '        for j in range(i + 1, n):',
    '            # Comparar fechas ISO "YYYY-MM-DD" como cadenas',
    '            if lista[j].get("date", "") > lista[max_idx].get("date", ""):',
    '                max_idx = j',
    '        if max_idx != i:',
    '            lista[i], lista[max_idx] = lista[max_idx], lista[i]',
    '    return lista',
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECCIÓN 6 – INSTRUCCIONES DE EJECUCIÓN
# ═══════════════════════════════════════════════════════════════
add_heading('6. Instrucciones de Ejecución')
add_divider()

add_heading('Modo 1 — Solo navegador (sin Python)', level=2)
add_body('Abre fintrack.html directamente en cualquier navegador moderno. '
         'Todas las funciones estarán disponibles excepto la exportación PDF.')

add_heading('Modo 2 — Con servidor Python', level=2)
add_code_block([
    '# 1. Instalar dependencia opcional para PDF',
    'pip install reportlab',
    '',
    '# 2. Colocar fintrack.html y fintrack_server.py en la misma carpeta',
    '',
    '# 3. Ejecutar el servidor',
    'python fintrack_server.py',
    '',
    '# 4. El navegador se abre automáticamente en:',
    '#    http://localhost:5000',
])

doc.add_paragraph()
add_heading('Verificación de funcionamiento', level=2)
pasos = [
    'Al abrir la app por primera vez, aparece una pantalla de bienvenida para configurar el saldo inicial.',
    'En el panel Dashboard se pueden agregar transacciones con el botón verde "Agregar".',
    'La tarjeta de saldo se actualiza en tiempo real al guardar cada transacción.',
    'En la vista Analytics aparece la gráfica donut con los gastos agrupados por categoría.',
    'El botón "Exportar CSV" descarga un archivo con todas las transacciones.',
    'Con el servidor Python activo, el botón PDF en Analytics genera el informe descargable.',
]
for paso in pasos:
    add_bullet(paso)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECCIÓN 7 – CONCLUSIONES
# ═══════════════════════════════════════════════════════════════
add_heading('7. Conclusiones')
add_divider()

conclusiones = [
    ('Aplicación del paradigma estructurado',
     'Se demostró que una aplicación web compleja y visualmente atractiva puede '
     'implementarse usando únicamente variables, arrays, funciones y estructuras '
     'de control básicas, sin necesidad de programación orientada a objetos.'),
    ('Trazabilidad de la lógica',
     'Al reemplazar objetos y métodos por arrays paralelos y funciones puras, '
     'cada paso del algoritmo queda visible y explícito. Esto facilita el '
     'aprendizaje y la revisión del código.'),
    ('Separación de responsabilidades',
     'Tanto en el frontend como en el backend se separaron claramente la '
     'presentación (renderizado), la lógica de negocio (cálculos) y el '
     'almacenamiento (persistencia), sin necesidad de clases.'),
    ('Uso responsable de la IA',
     'Claude (Anthropic) fue utilizado como herramienta de revisión y '
     'documentación, no como sustituto del proceso de diseño. '
     'Las decisiones de arquitectura y la comprensión del código '
     'son responsabilidad del equipo de desarrollo.'),
]

for titulo, texto in conclusiones:
    add_bullet(texto, bold_part=titulo)

doc.add_paragraph()
add_divider()
p_fin = doc.add_paragraph()
p_fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_fin = p_fin.add_run('© 2026 FinTrack — Documentación Técnica Educativa')
r_fin.font.size = Pt(9); r_fin.font.color.rgb = GRAY; r_fin.font.name = 'Calibri'

# ── Guardar ──────────────────────────────────────────────────────────────────
doc.save('/mnt/user-data/outputs/FinTrack_Documentacion_Tecnica.docx')
print("Documento generado correctamente.")